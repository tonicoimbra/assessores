"""Document extraction adapter supporting PDF and DOCX."""

from __future__ import annotations

import logging
from hashlib import sha256
from pathlib import Path

from docx import Document as DocxDocument

from src.models import DocumentoEntrada, TipoDocumento
from src.pdf_processor import (
    PDFExtractionError,
    ExtractionResult,
    _calcular_noise_ratio,
    _calcular_score_qualidade,
    _limpar_texto,
    extrair_texto as extrair_texto_pdf,
)

logger = logging.getLogger("assessor_ai")

SUPPORTED_DOCUMENT_EXTENSIONS: set[str] = {".pdf", ".docx"}


class DocumentExtractionError(PDFExtractionError):
    """Raised when generic document extraction fails."""


def _normalizar_docx_texto(document: DocxDocument) -> str:
    """Collect relevant textual content from paragraphs and tables."""
    blocos: list[str] = []

    for paragraph in document.paragraphs:
        texto = str(paragraph.text or "").strip()
        if texto:
            blocos.append(texto)

    for table in document.tables:
        for row in table.rows:
            valores = [str(cell.text or "").strip() for cell in row.cells]
            linha = " | ".join(valor for valor in valores if valor)
            if linha:
                blocos.append(linha)

    return "\n".join(blocos).strip()


def extrair_texto_docx(filepath: str) -> ExtractionResult:
    """Extract text from a DOCX file into ExtractionResult."""
    path = Path(filepath)
    if not path.exists():
        raise DocumentExtractionError(f"Arquivo não encontrado: {filepath}")

    if path.suffix.lower() != ".docx":
        raise DocumentExtractionError(
            f"Formato inválido para DOCX extractor: {path.suffix}"
        )

    try:
        document = DocxDocument(filepath)
    except Exception as exc:  # pragma: no cover - defensive branch
        raise DocumentExtractionError(
            f"Não foi possível extrair texto de {path.name}: {exc}"
        ) from exc

    raw_text = _normalizar_docx_texto(document)
    clean_text = _limpar_texto(raw_text)
    pagina_unica = raw_text if raw_text else ""
    num_paginas = 1 if pagina_unica else 0

    resultado = ExtractionResult(
        texto=clean_text,
        num_paginas=num_paginas,
        num_caracteres=len(clean_text),
        engine_usada="docx",
        raw_text_by_page=[pagina_unica] if pagina_unica else [],
        clean_text_by_page=[clean_text] if clean_text else [],
        raw_page_hashes=[sha256(pagina_unica.encode("utf-8")).hexdigest()] if pagina_unica else [],
        clean_page_hashes=[sha256(clean_text.encode("utf-8")).hexdigest()] if clean_text else [],
        quality_score_by_page=[_calcular_score_qualidade(clean_text)] if clean_text else [],
        noise_ratio_by_page=[_calcular_noise_ratio(pagina_unica)] if pagina_unica else [],
        quality_score=_calcular_score_qualidade(clean_text) if clean_text else 0.0,
        noise_ratio=_calcular_noise_ratio(pagina_unica) if pagina_unica else 1.0,
        ocr_confidence_by_page=[0.0] if pagina_unica else [],
        ocr_confidence=0.0,
        ocr_aplicado=False,
    )
    return resultado


def extract_text(filepath: str) -> ExtractionResult:
    """Extract text from supported document formats (.pdf, .docx)."""
    path = Path(filepath)
    if not path.exists():
        raise DocumentExtractionError(f"Arquivo não encontrado: {filepath}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extrair_texto_pdf(filepath)
    if suffix == ".docx":
        return extrair_texto_docx(filepath)

    raise DocumentExtractionError(
        f"Formato inválido (esperado .pdf ou .docx): {path.suffix}"
    )


def extrair_multiplos_documentos(filepaths: list[str]) -> list[DocumentoEntrada]:
    """Extract text from a list of documents supporting PDF and DOCX."""
    documentos: list[DocumentoEntrada] = []

    for filepath in filepaths:
        try:
            resultado = extract_text(filepath)
        except PDFExtractionError as exc:
            logger.error("❌ Erro ao processar %s: %s", filepath, exc)
            raise

        documentos.append(
            DocumentoEntrada(
                filepath=filepath,
                texto_extraido=resultado.texto,
                tipo=TipoDocumento.DESCONHECIDO,
                num_paginas=resultado.num_paginas,
                num_caracteres=resultado.num_caracteres,
            )
        )

    return documentos

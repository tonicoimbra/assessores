"""Output formatter: markdown formatting, file saving, and audit reports."""

import json
import logging
import re
from hashlib import sha256
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:  # pragma: no cover - covered by runtime guard
    Document = None  # type: ignore[assignment]
    WD_PARAGRAPH_ALIGNMENT = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    Cm = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]

from src.config import OUTPUTS_DIR
from src.models import EstadoPipeline, ResultadoEtapa3

logger = logging.getLogger("assessor_ai")

INLINE_MARKDOWN_RE = re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)")


def _resolver_output_dir(output_dir: Path | None = None) -> Path:
    """Resolve and ensure output directory exists."""
    target_dir = output_dir or OUTPUTS_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    return target_dir


# --- 5.3.1 / 5.3.2 Format draft ---


def formatar_minuta(resultado: ResultadoEtapa3, estado: EstadoPipeline | None = None) -> str:
    """Format the admissibility draft with proper markdown styling."""
    minuta = resultado.minuta_completa

    # 5.3.2 — Bold mandatory fields
    if estado and estado.resultado_etapa1:
        r1 = estado.resultado_etapa1
        if r1.recorrente:
            minuta = minuta.replace(r1.recorrente, f"**{r1.recorrente}**")
        if r1.recorrido:
            minuta = minuta.replace(r1.recorrido, f"**{r1.recorrido}**")
        if r1.especie_recurso:
            minuta = minuta.replace(r1.especie_recurso, f"**{r1.especie_recurso}**")

    # Bold the decision keywords
    for pattern, repl in [
        (r"\b(ADMITO)\b", r"**\1**"),
        (r"\b(INADMITO)\b", r"**\1**"),
        (r"\b(INCONCLUSIVO)\b", r"**\1**"),
    ]:
        minuta = re.sub(pattern, repl, minuta)

    # Avoid double-bold
    minuta = minuta.replace("****", "**")

    return minuta


# --- 5.3.3 Save draft ---


def salvar_minuta(
    minuta_formatada: str,
    numero_processo: str = "sem_numero",
    output_dir: Path | None = None,
) -> Path:
    """Save formatted draft to markdown file in outputs/ directory."""
    target_dir = _resolver_output_dir(output_dir)

    safe_proc = re.sub(r"[^\w\-.]", "_", numero_processo)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"minuta_{safe_proc}_{timestamp}.md"
    filepath = target_dir / filename

    filepath.write_text(minuta_formatada, encoding="utf-8")
    logger.info("📄 Minuta salva em: %s", filepath)

    return filepath


def _add_markdown_runs(paragraph: "object", text: str) -> None:
    """Add simple markdown inline formatting (bold/italic) to a DOCX paragraph."""
    for token in INLINE_MARKDOWN_RE.split(text):
        if not token:
            continue
        is_bold = (
            token.startswith("**") and token.endswith("**")
        ) or (token.startswith("__") and token.endswith("__"))
        is_italic = (
            token.startswith("*") and token.endswith("*")
        ) or (token.startswith("_") and token.endswith("_"))

        clean = token
        if is_bold or is_italic:
            clean = token[2:-2] if len(token) > 4 and token[:2] in {"**", "__"} else token[1:-1]

        run = paragraph.add_run(clean)
        run.bold = is_bold
        run.italic = is_italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _configure_docx_tjpr(document: "object") -> None:
    """Apply TJPR-compatible baseline formatting to DOCX document."""
    section = document.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.0)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)


def _add_paragraph_safe(document: "object", style: str | None = None) -> "object":
    """Add paragraph with optional style fallback when style is unavailable."""
    if style:
        try:
            return document.add_paragraph(style=style)
        except KeyError:
            logger.debug("Style '%s' não encontrado no template DOCX. Usando estilo padrão.", style)
    return document.add_paragraph()


def _apply_tjpr_paragraph_format(paragraph: "object", justify: bool = True) -> None:
    """Apply paragraph formatting rules used in TJPR-compatible output."""
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if justify:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def salvar_minuta_docx(
    minuta_markdown: str,
    numero_processo: str = "sem_numero",
    output_dir: Path | None = None,
) -> Path:
    """Convert markdown draft to DOCX and apply TJPR-compatible formatting."""
    if Document is None:
        raise RuntimeError(
            "Dependência opcional 'python-docx' não instalada. "
            "Adicione no ambiente para usar --formato docx."
        )

    target_dir = _resolver_output_dir(output_dir)
    safe_proc = re.sub(r"[^\w\-.]", "_", numero_processo)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = target_dir / f"minuta_{safe_proc}_{timestamp}.docx"

    document = Document()
    _configure_docx_tjpr(document)

    in_code_block = False
    for raw_line in minuta_markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if not stripped:
            document.add_paragraph("")
            continue

        if in_code_block:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            paragraph.paragraph_format.line_spacing = 1.0
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = heading_match.group(2).strip()
            paragraph = document.add_heading("", level=level)
            _apply_tjpr_paragraph_format(paragraph, justify=False)
            _add_markdown_runs(paragraph, text)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            paragraph = _add_paragraph_safe(document, style="List Bullet")
            _apply_tjpr_paragraph_format(paragraph)
            _add_markdown_runs(paragraph, bullet_match.group(1))
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            paragraph = _add_paragraph_safe(document, style="List Number")
            _apply_tjpr_paragraph_format(paragraph)
            _add_markdown_runs(paragraph, numbered_match.group(1))
            continue

        paragraph = document.add_paragraph()
        _apply_tjpr_paragraph_format(paragraph)
        _add_markdown_runs(paragraph, stripped)

    document.save(filepath)
    logger.info("📄 Minuta DOCX salva em: %s", filepath)
    return filepath


# --- 5.3.4 Audit report ---


def _build_audit_payload(
    estado: EstadoPipeline,
    alertas: list[str] | None = None,
    numero_processo: str = "sem_numero",
) -> dict:
    """Build structured audit payload shared by markdown and JSON outputs."""
    meta = estado.metadata
    r1 = estado.resultado_etapa1
    r2 = estado.resultado_etapa2
    r3 = estado.resultado_etapa3

    return {
        "processo": numero_processo,
        "execucao_id": meta.execucao_id,
        "gerado_em": datetime.now().isoformat(),
        "tokens": {
            "modelo": meta.modelo_usado or "N/A",
            "prompt_tokens": meta.prompt_tokens,
            "completion_tokens": meta.completion_tokens,
            "total_tokens": meta.total_tokens,
        },
        "llm_stats": meta.llm_stats,
        "prompt": {
            "profile": meta.prompt_profile,
            "version": meta.prompt_version,
            "hash_sha256": meta.prompt_hash_sha256,
        },
        "modelos_utilizados": {
            "etapa1": meta.modelos_utilizados.get("Etapa 1", "N/A"),
            "etapa2": meta.modelos_utilizados.get("Etapa 2", "N/A"),
            "etapa3": meta.modelos_utilizados.get("Etapa 3", "N/A"),
        },
        "confianca": {
            "por_etapa": meta.confianca_por_etapa,
            "por_campo_etapa1": meta.confianca_campos_etapa1,
            "por_tema_etapa2": meta.confianca_temas_etapa2,
            "global": meta.confianca_global,
        },
        "escalonamento": meta.politica_escalonamento,
        "classificacao_revisao_manual": meta.classificacao_revisao_manual,
        "chunking_auditoria": meta.chunking_auditoria,
        "motivo_bloqueio": {
            "codigo": meta.motivo_bloqueio_codigo,
            "descricao": meta.motivo_bloqueio_descricao,
        },
        "pipeline": {
            "etapa1": bool(r1),
            "etapa2": {
                "executada": bool(r2),
                "temas": len(r2.temas) if r2 else 0,
            },
            "etapa3": {
                "executada": bool(r3),
                "decisao": r3.decisao.value if r3 and r3.decisao else "N/A",
                "fundamentos_decisao": r3.fundamentos_decisao if r3 else [],
                "itens_evidencia_usados": r3.itens_evidencia_usados if r3 else [],
                "aviso_inconclusivo": bool(r3.aviso_inconclusivo) if r3 else False,
                "motivo_bloqueio_codigo": r3.motivo_bloqueio_codigo if r3 else "",
                "motivo_bloqueio_descricao": r3.motivo_bloqueio_descricao if r3 else "",
            },
        },
        "timeline": {
            "inicio": meta.inicio.isoformat() if meta.inicio else None,
            "fim": meta.fim.isoformat() if meta.fim else None,
        },
        "alertas_validacao": alertas or [],
    }


def gerar_relatorio_auditoria(
    estado: EstadoPipeline,
    alertas: list[str] | None = None,
    numero_processo: str = "sem_numero",
    output_dir: Path | None = None,
) -> Path:
    """Generate audit report alongside the draft."""
    target_dir = _resolver_output_dir(output_dir)
    safe_proc = re.sub(r"[^\w\-.]", "_", numero_processo)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = target_dir / f"auditoria_{safe_proc}_{timestamp}.md"
    payload = _build_audit_payload(estado, alertas=alertas, numero_processo=numero_processo)

    tokens = payload["tokens"]
    llm_stats = payload["llm_stats"]
    prompt = payload["prompt"]
    modelos = payload["modelos_utilizados"]
    confianca = payload["confianca"]
    escalonamento = payload["escalonamento"]
    classificacao_revisao_manual = payload["classificacao_revisao_manual"]
    chunking_auditoria = payload["chunking_auditoria"]
    motivo_bloqueio = payload["motivo_bloqueio"]
    pipeline = payload["pipeline"]
    timeline = payload["timeline"]
    alertas_lista = payload["alertas_validacao"]

    lines = [
        "# Relatório de Auditoria",
        "",
        f"**Processo:** {numero_processo}",
        f"**Execução:** {payload['execucao_id'] or 'N/A'}",
        f"**Data:** {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}",
        "",
        "## Tokens Utilizados",
        "",
        "| Métrica | Valor |",
        "|---------|-------|",
        f"| Modelo | {tokens['modelo']} |",
        f"| Prompt tokens | {tokens['prompt_tokens']} |",
        f"| Completion tokens | {tokens['completion_tokens']} |",
        f"| **Total tokens** | **{tokens['total_tokens']}** |",
        "",
        "## Métricas LLM",
        "",
        f"- Chamadas totais: {int(llm_stats.get('total_calls', 0))}",
        f"- Chamadas truncadas: {int(llm_stats.get('calls_truncadas', 0))}",
        f"- Latência média: {llm_stats.get('latencia_media_ms', 0.0):.2f} ms",
        "",
        "## Assinatura de Prompt",
        "",
        f"- Profile: `{prompt['profile'] or 'N/A'}`",
        f"- Version: `{prompt['version'] or 'N/A'}`",
        f"- Hash SHA-256: `{prompt['hash_sha256'] or 'N/A'}`",
        "",
        "## Modelos Utilizados",
        "",
        "| Etapa | Modelo |",
        "|-------|--------|",
        f"| Etapa 1 (Análise Recursal) | {modelos['etapa1']} |",
        f"| Etapa 2 (Análise Direito) | {modelos['etapa2']} |",
        f"| Etapa 3 (Geração Minuta) | {modelos['etapa3']} |",
        "",
        "## Confiança",
        "",
        f"- Global: **{confianca['global']:.3f}**",
        f"- Etapa 1: {confianca['por_etapa'].get('etapa1', 0.0):.3f}",
        f"- Etapa 2: {confianca['por_etapa'].get('etapa2', 0.0):.3f}",
        f"- Etapa 3: {confianca['por_etapa'].get('etapa3', 0.0):.3f}",
        f"- Campos críticos (Etapa 1): {confianca['por_campo_etapa1'] or {}}",
        f"- Temas (Etapa 2): {confianca['por_tema_etapa2'] or {}}",
        "",
        "## Escalonamento por Confiança",
        "",
        f"- Ativo: {'Sim' if bool(escalonamento.get('ativo')) else 'Não'}",
        f"- Escalonar revisão humana: {'Sim' if bool(escalonamento.get('escalonar')) else 'Não'}",
        f"- Thresholds: {escalonamento.get('thresholds', {})}",
        "",
        "## Revisão Manual de Classificação",
        "",
        f"- Ativo: {'Sim' if bool(classificacao_revisao_manual.get('ativo')) else 'Não'}",
        f"- Revisão recomendada: {'Sim' if bool(classificacao_revisao_manual.get('revisao_recomendada')) else 'Não'}",
        f"- Thresholds: {classificacao_revisao_manual.get('thresholds', {})}",
        f"- Documentos ambíguos: {classificacao_revisao_manual.get('documentos_ambiguos', [])}",
        "",
        "## Chunking Auditável",
        "",
        f"- Dados de chunking por etapa: {chunking_auditoria or {}}",
        "",
        "## Motivo de Bloqueio",
        "",
        f"- Código: {motivo_bloqueio['codigo'] or 'N/A'}",
        f"- Descrição: {motivo_bloqueio['descricao'] or 'N/A'}",
        "",
        "## Pipeline",
        "",
        "| Etapa | Status |",
        "|-------|--------|",
        f"| Etapa 1 | {'✅' if pipeline['etapa1'] else '❌'} |",
        f"| Etapa 2 | {'✅ ' + str(pipeline['etapa2']['temas']) + ' temas' if pipeline['etapa2']['executada'] else '❌'} |",
        f"| Etapa 3 | {'✅ ' + pipeline['etapa3']['decisao'] if pipeline['etapa3']['executada'] else '❌'} |",
        "",
    ]

    if timeline["inicio"]:
        lines.append(f"**Início:** {timeline['inicio']}")
    if timeline["fim"]:
        lines.append(f"**Fim:** {timeline['fim']}")

    if alertas_lista:
        lines.extend(["", "## Alertas de Validação", ""])
        for alerta in alertas_lista:
            lines.append(f"- ⚠️ {alerta}")
    motivos_escalonamento = escalonamento.get("motivos", [])
    if motivos_escalonamento:
        lines.extend(["", "## Motivos de Escalonamento", ""])
        for motivo in motivos_escalonamento:
            lines.append(f"- {motivo}")

    lines.append("")

    filepath.write_text("\n".join(lines), encoding="utf-8")
    logger.info("📊 Relatório de auditoria salvo em: %s", filepath)

    return filepath


def salvar_trilha_auditoria_json(
    estado: EstadoPipeline,
    alertas: list[str] | None = None,
    numero_processo: str = "sem_numero",
    output_dir: Path | None = None,
) -> Path:
    """Save structured audit trail as JSON."""
    target_dir = _resolver_output_dir(output_dir)
    safe_proc = re.sub(r"[^\w\-.]", "_", numero_processo)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = target_dir / f"auditoria_{safe_proc}_{timestamp}.json"

    payload = _build_audit_payload(estado, alertas=alertas, numero_processo=numero_processo)
    filepath.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    logger.info("🧾 Trilha de auditoria JSON salva em: %s", filepath)
    return filepath


def _preview_text(texto: str, limit: int = 500) -> str:
    """Return compact preview of a potentially large text."""
    bruto = str(texto or "")
    compact = re.sub(r"\s+", " ", bruto).strip()
    if len(compact) <= limit:
        return compact
    return compact[:limit].rstrip() + "..."


def _hash_text(texto: str) -> str:
    """Return stable SHA-256 hash for source text traceability."""
    return sha256(str(texto or "").encode("utf-8")).hexdigest()


def salvar_snapshot_execucao_json(
    estado: EstadoPipeline,
    validacoes: dict[str, list[str]] | None = None,
    arquivos_saida: dict[str, str] | None = None,
    numero_processo: str = "sem_numero",
    output_dir: Path | None = None,
) -> Path:
    """Persist full execution snapshot with inputs, outputs and stage validations."""
    target_dir = _resolver_output_dir(output_dir)
    safe_proc = re.sub(r"[^\w\-.]", "_", numero_processo)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = target_dir / f"snapshot_execucao_{safe_proc}_{timestamp}.json"

    documentos = []
    for doc in estado.documentos_entrada:
        documentos.append(
            {
                "filepath": doc.filepath,
                "tipo": doc.tipo.value if doc.tipo else "DESCONHECIDO",
                "num_paginas": doc.num_paginas,
                "num_caracteres": doc.num_caracteres,
                "texto_extraido_hash": _hash_text(doc.texto_extraido),
                "texto_extraido_preview": _preview_text(doc.texto_extraido),
            }
        )

    payload = {
        "snapshot_schema_version": "1.0.0",
        "processo_id": numero_processo,
        "gerado_em": datetime.now().isoformat(),
        "inputs": {
            "documentos": documentos,
        },
        "stages": {
            "etapa1": {
                "executada": estado.resultado_etapa1 is not None,
                "resultado": (
                    estado.resultado_etapa1.model_dump(mode="json")
                    if estado.resultado_etapa1
                    else None
                ),
                "validacao_erros": (validacoes or {}).get("etapa1", []),
            },
            "etapa2": {
                "executada": estado.resultado_etapa2 is not None,
                "resultado": (
                    estado.resultado_etapa2.model_dump(mode="json")
                    if estado.resultado_etapa2
                    else None
                ),
                "validacao_erros": (validacoes or {}).get("etapa2", []),
            },
            "etapa3": {
                "executada": estado.resultado_etapa3 is not None,
                "resultado": (
                    estado.resultado_etapa3.model_dump(mode="json")
                    if estado.resultado_etapa3
                    else None
                ),
                "validacao_erros": (validacoes or {}).get("etapa3", []),
            },
        },
        "metadata": estado.metadata.model_dump(mode="json"),
        "outputs": arquivos_saida or {},
    }

    filepath.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    logger.info("🧷 Snapshot de execução salvo em: %s", filepath)
    return filepath

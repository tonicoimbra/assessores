"""Tests for generic document extraction (.pdf/.docx)."""

from pathlib import Path

from docx import Document as DocxDocument

import pytest

from src.document_extractor import (
    DocumentExtractionError,
    extract_text,
    extrair_multiplos_documentos,
)
from src.pdf_processor import ExtractionResult


def test_extract_text_docx_returns_extraction_result(tmp_path: Path) -> None:
    docx_path = tmp_path / "peticao.docx"
    doc = DocxDocument()
    doc.add_paragraph("Recurso especial interposto com fundamento no art. 105, III.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Tema"
    table.rows[0].cells[1].text = "Responsabilidade civil"
    doc.save(docx_path)

    resultado = extract_text(str(docx_path))

    assert resultado.engine_usada == "docx"
    assert resultado.num_paginas == 1
    assert resultado.num_caracteres > 0
    assert "Recurso especial interposto" in resultado.texto
    assert "Tema | Responsabilidade civil" in resultado.texto


def test_extract_text_dispatches_pdf_to_pdf_adapter(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pdf_path = tmp_path / "entrada.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n")

    monkeypatch.setattr(
        "src.document_extractor.extrair_texto_pdf",
        lambda _filepath: ExtractionResult(
            texto="texto-pdf",
            num_paginas=2,
            num_caracteres=9,
            engine_usada="pymupdf",
        ),
    )

    resultado = extract_text(str(pdf_path))
    assert resultado.engine_usada == "pymupdf"
    assert resultado.texto == "texto-pdf"


def test_extract_text_rejects_invalid_extension(tmp_path: Path) -> None:
    invalid_path = tmp_path / "entrada.txt"
    invalid_path.write_text("conteudo", encoding="utf-8")

    with pytest.raises(DocumentExtractionError, match="Formato inválido"):
        extract_text(str(invalid_path))


def test_extrair_multiplos_documentos_supports_pdf_and_docx(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fake_result = ExtractionResult(
        texto="texto extraido",
        num_paginas=1,
        num_caracteres=13,
        engine_usada="docx",
    )
    monkeypatch.setattr("src.document_extractor.extract_text", lambda _filepath: fake_result)

    documentos = extrair_multiplos_documentos(["/tmp/recurso.pdf", "/tmp/acordao.docx"])
    assert len(documentos) == 2
    assert documentos[0].filepath.endswith(".pdf")
    assert documentos[1].filepath.endswith(".docx")
    assert all(doc.texto_extraido == "texto extraido" for doc in documentos)


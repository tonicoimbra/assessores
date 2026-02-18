"""Tests for Sprint 5: Etapa 3 validation, cross-checking, and output formatting."""

import shutil
from importlib.util import find_spec
from pathlib import Path
from unittest.mock import patch

import pytest

from src.etapa3 import (
    Etapa3Error,
    _decidir_admissibilidade_deterministica,
    _extrair_decisao,
    _resultado_etapa3_from_json,
    _validar_cruzada_dispositivos,
    _validar_cruzada_temas,
    _validar_secao_i,
    _validar_secao_ii,
    _validar_secao_iii,
    _validar_secoes,
    _validar_sumulas_secao_iii,
    _validar_transcricoes,
    executar_etapa3,
)
from src.models import (
    DocumentoEntrada,
    CampoEvidencia,
    Decisao,
    EstadoPipeline,
    MetadadosPipeline,
    ResultadoEtapa1,
    ResultadoEtapa2,
    ResultadoEtapa3,
    TemaEtapa2,
    TipoDocumento,
)
from src.output_formatter import (
    formatar_minuta,
    gerar_relatorio_auditoria,
    salvar_snapshot_execucao_json,
    salvar_minuta,
    salvar_trilha_auditoria_json,
    salvar_minuta_docx,
)

HAS_PYTHON_DOCX = find_spec("docx") is not None


MINUTA_COMPLETA = """
Seção I – Relatório
Processo nº 0001234-56.2024.8.16.0001
Recorrente: EMPRESA ABC LTDA
Recorrido: CONSUMIDOR XYZ
Espécie: RECURSO ESPECIAL
Dispositivos violados: art. 927 do CC

Seção II – Análise Temática
Tema: Responsabilidade civil
O tribunal entendeu que não há dano moral comprovado.
"Não restou demonstrada a ocorrência do dano alegado pelo recorrente."

Seção III – Decisão
Ante o exposto, INADMITO o recurso especial interposto.
Fundamento: Súmula 7 do STJ.
"""


# --- 5.5.1: Structure I/II/III validation ---


class TestValidacaoEstrutura:
    """Test validation of draft sections I, II, III."""

    def test_valid_structure_no_alerts(self) -> None:
        alertas = _validar_secoes(MINUTA_COMPLETA)
        assert len(alertas) == 0

    def test_missing_section_i(self) -> None:
        sem_i = MINUTA_COMPLETA.replace("Seção I", "Parte X")
        alertas = _validar_secoes(sem_i)
        assert any("I" in a for a in alertas)

    def test_missing_section_iii(self) -> None:
        sem_iii = MINUTA_COMPLETA.replace("Seção III", "Conclusão")
        alertas = _validar_secoes(sem_iii)
        assert any("III" in a for a in alertas)

    def test_detects_admito(self) -> None:
        assert _extrair_decisao("ADMITO o recurso") == Decisao.ADMITIDO

    def test_detects_inadmito(self) -> None:
        assert _extrair_decisao(MINUTA_COMPLETA) == Decisao.INADMITIDO


# --- 5.5.2: Divergence Etapa 1 x Seção I ---


class TestDivergenciaEtapa1:
    """Test cross-check between Stage 1 and section I."""

    def test_no_divergence(self) -> None:
        r1 = ResultadoEtapa1(
            numero_processo="0001234-56.2024.8.16.0001",
            recorrente="EMPRESA ABC LTDA",
            especie_recurso="RECURSO ESPECIAL",
        )
        alertas = _validar_secao_i(MINUTA_COMPLETA, r1)
        assert len(alertas) == 0

    def test_missing_process_number(self) -> None:
        r1 = ResultadoEtapa1(numero_processo="9999999-99.9999.9.99.9999")
        alertas = _validar_secao_i(MINUTA_COMPLETA, r1)
        assert any("número" in a for a in alertas)

    def test_missing_recorrente(self) -> None:
        r1 = ResultadoEtapa1(recorrente="PESSOA INEXISTENTE")
        alertas = _validar_secao_i(MINUTA_COMPLETA, r1)
        assert any("recorrente" in a.lower() for a in alertas)

    def test_cross_validate_devices(self) -> None:
        r1 = ResultadoEtapa1(dispositivos_violados=["art. 927 do CC"])
        alertas = _validar_cruzada_dispositivos(MINUTA_COMPLETA, r1)
        assert len(alertas) == 0

    def test_cross_validate_missing_device(self) -> None:
        r1 = ResultadoEtapa1(dispositivos_violados=["art. 999 do CPC"])
        alertas = _validar_cruzada_dispositivos(MINUTA_COMPLETA, r1)
        assert any("999" in a for a in alertas)


# --- 5.5.3: Transcription in ruling ---


class TestTranscricaoLiteral:
    """Test that transcription excerpts exist in ruling text."""

    def test_valid_transcription(self) -> None:
        acordao = "Não restou demonstrada a ocorrência do dano alegado pelo recorrente."
        alertas = _validar_transcricoes(MINUTA_COMPLETA, acordao)
        assert len(alertas) == 0

    def test_missing_transcription(self) -> None:
        alertas = _validar_transcricoes(MINUTA_COMPLETA, "Texto completamente diferente.")
        assert any("Transcrição" in a for a in alertas)

    def test_sumula_consistency(self) -> None:
        r2 = ResultadoEtapa2(temas=[TemaEtapa2(obices_sumulas=["Súmula 7"])])
        alertas = _validar_sumulas_secao_iii(MINUTA_COMPLETA, r2)
        assert len(alertas) == 0

    def test_sumula_new_in_draft(self) -> None:
        r2 = ResultadoEtapa2(temas=[])
        alertas = _validar_sumulas_secao_iii(MINUTA_COMPLETA, r2)
        assert any("7" in a for a in alertas)


# --- 5.5.4: Markdown formatting ---


class TestFormatacaoMarkdown:
    """Test markdown formatting of the draft."""

    def test_bold_recorrente(self) -> None:
        r3 = ResultadoEtapa3(minuta_completa="Recorrente: FULANO de TAL")
        r1 = ResultadoEtapa1(recorrente="FULANO de TAL")
        estado = EstadoPipeline(resultado_etapa1=r1)
        fmt = formatar_minuta(r3, estado)
        assert "**FULANO de TAL**" in fmt

    def test_bold_decisao(self) -> None:
        r3 = ResultadoEtapa3(minuta_completa="INADMITO o recurso.")
        fmt = formatar_minuta(r3)
        assert "**INADMITO**" in fmt

    def test_no_double_bold(self) -> None:
        r3 = ResultadoEtapa3(minuta_completa="**ADMITO** o recurso.")
        fmt = formatar_minuta(r3)
        assert "****" not in fmt

    def test_save_minuta_creates_file(self, tmp_path: Path) -> None:
        with patch("src.output_formatter.OUTPUTS_DIR", tmp_path):
            path = salvar_minuta("# Minuta\nConteúdo.", "123")
            assert path.exists()
            assert "minuta_123" in path.name
            assert path.read_text(encoding="utf-8").startswith("# Minuta")

    def test_audit_report(self, tmp_path: Path) -> None:
        with patch("src.output_formatter.OUTPUTS_DIR", tmp_path):
            estado = EstadoPipeline(
                resultado_etapa1=ResultadoEtapa1(numero_processo="123"),
                resultado_etapa2=ResultadoEtapa2(temas=[TemaEtapa2()]),
                resultado_etapa3=ResultadoEtapa3(decisao=Decisao.INADMITIDO),
                metadata=MetadadosPipeline(
                    total_tokens=5000,
                    modelo_usado="gpt-4o",
                    execucao_id="exec-123",
                    prompt_profile="modular-lean",
                    prompt_version="unversioned",
                    prompt_hash_sha256="a" * 64,
                    llm_stats={"total_calls": 3, "calls_truncadas": 1, "latencia_media_ms": 120.5},
                ),
            )
            path = gerar_relatorio_auditoria(estado, ["teste"], "123")
            content = path.read_text(encoding="utf-8")
            assert "5000" in content
            assert "teste" in content
            assert "gpt-4o" in content
            assert "exec-123" in content
            assert "Chamadas truncadas: 1" in content
            assert "Assinatura de Prompt" in content
            assert "modular-lean" in content

    def test_audit_json_report(self, tmp_path: Path) -> None:
        with patch("src.output_formatter.OUTPUTS_DIR", tmp_path):
            estado = EstadoPipeline(
                resultado_etapa1=ResultadoEtapa1(numero_processo="123"),
                resultado_etapa2=ResultadoEtapa2(temas=[TemaEtapa2()]),
                resultado_etapa3=ResultadoEtapa3(
                    decisao=Decisao.INCONCLUSIVO,
                    fundamentos_decisao=["Dados insuficientes"],
                    itens_evidencia_usados=["Tema 1 sem evidência suficiente"],
                    aviso_inconclusivo=True,
                    motivo_bloqueio_codigo="E3_INCONCLUSIVO",
                    motivo_bloqueio_descricao="Dados insuficientes para decisão conclusiva.",
                ),
                metadata=MetadadosPipeline(
                    total_tokens=5000,
                    modelo_usado="gpt-4o",
                    execucao_id="exec-123",
                    prompt_profile="modular-lean",
                    prompt_version="unversioned",
                    prompt_hash_sha256="b" * 64,
                    llm_stats={"total_calls": 3, "calls_truncadas": 1, "latencia_media_ms": 120.5},
                    confianca_por_etapa={"etapa1": 0.91, "etapa2": 0.77, "etapa3": 0.49},
                    confianca_global=0.72,
                    motivo_bloqueio_codigo="E3_INCONCLUSIVO",
                    motivo_bloqueio_descricao="Dados insuficientes para decisão conclusiva.",
                ),
            )
            path = salvar_trilha_auditoria_json(estado, ["teste"], "123")
            payload = path.read_text(encoding="utf-8")
            assert path.suffix == ".json"
            assert "\"processo\": \"123\"" in payload
            assert "\"execucao_id\": \"exec-123\"" in payload
            assert "\"decisao\": \"INCONCLUSIVO\"" in payload
            assert "\"alertas_validacao\"" in payload
            assert "\"confianca\"" in payload
            assert "\"global\": 0.72" in payload
            assert "\"escalonamento\"" in payload
            assert "\"classificacao_revisao_manual\"" in payload
            assert "\"chunking_auditoria\"" in payload
            assert "\"llm_stats\"" in payload
            assert "\"prompt\"" in payload
            assert "\"profile\": \"modular-lean\"" in payload
            assert "\"motivo_bloqueio\"" in payload
            assert "\"codigo\": \"E3_INCONCLUSIVO\"" in payload

    def test_execution_snapshot_json(self, tmp_path: Path) -> None:
        with patch("src.output_formatter.OUTPUTS_DIR", tmp_path):
            estado = EstadoPipeline(
                documentos_entrada=[
                    DocumentoEntrada(
                        filepath="recurso.pdf",
                        tipo=TipoDocumento.RECURSO,
                        num_paginas=3,
                        num_caracteres=1200,
                        texto_extraido="Recurso especial com fundamentos e pedidos.",
                    ),
                ],
                resultado_etapa1=ResultadoEtapa1(numero_processo="123", recorrente="EMPRESA X"),
                resultado_etapa2=ResultadoEtapa2(temas=[TemaEtapa2(materia_controvertida="Tema A")]),
                resultado_etapa3=ResultadoEtapa3(decisao=Decisao.ADMITIDO),
                metadata=MetadadosPipeline(total_tokens=5000, modelo_usado="gpt-4o"),
            )
            path = salvar_snapshot_execucao_json(
                estado,
                validacoes={"etapa1": [], "etapa2": [], "etapa3": []},
                arquivos_saida={"minuta": "/tmp/minuta.md"},
                numero_processo="123",
            )
            payload = path.read_text(encoding="utf-8")
            assert path.suffix == ".json"
            assert "\"snapshot_schema_version\": \"1.0.0\"" in payload
            assert "\"processo_id\": \"123\"" in payload
            assert "\"texto_extraido_hash\"" in payload
            assert "\"validacao_erros\"" in payload
            assert "\"metadata\"" in payload

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx não instalado")
    def test_save_minuta_docx_creates_file(self, tmp_path: Path) -> None:
        from docx import Document

        with patch("src.output_formatter.OUTPUTS_DIR", tmp_path):
            path = salvar_minuta_docx("# Minuta\nTexto com **destaque**.", "123")
            assert path.exists()
            assert path.suffix == ".docx"
            doc = Document(path)
            assert any("Minuta" in p.text for p in doc.paragraphs)
            assert any(run.bold for p in doc.paragraphs for run in p.runs)

    @pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx não instalado")
    def test_docx_tjpr_formatting(self, tmp_path: Path) -> None:
        from docx import Document

        with patch("src.output_formatter.OUTPUTS_DIR", tmp_path):
            path = salvar_minuta_docx("Seção I\nParágrafo de teste.", "123")
            doc = Document(path)
            section = doc.sections[0]
            normal = doc.styles["Normal"]

            assert round(section.left_margin.cm, 1) == 3.0
            assert round(section.right_margin.cm, 1) == 2.0
            assert round(section.top_margin.cm, 1) == 3.0
            assert round(section.bottom_margin.cm, 1) == 2.0
            assert normal.font.name == "Times New Roman"
            assert int(normal.font.size.pt) == 12


# --- 5.5.5: Integration (slow) ---


@pytest.mark.slow
class TestPipelineIntegration:
    """Integration test: full pipeline Etapa 1→2→3. Run with: pytest -m slow"""

    def test_full_pipeline(self) -> None:
        from src.config import OPENAI_API_KEY

        if not OPENAI_API_KEY:
            pytest.skip("OPENAI_API_KEY not set")

        from src.etapa1 import executar_etapa1
        from src.etapa2 import executar_etapa2
        from src.etapa3 import executar_etapa3

        texto_recurso = "RECURSO ESPECIAL. Recorrente: TESTE. Processo 123."
        texto_acordao = "ACÓRDÃO. Vistos. EMENTA: Dano moral."

        r1 = executar_etapa1(texto_recurso, "Extraia dados do recurso.")
        r2 = executar_etapa2(texto_acordao, r1, "Analise temas do acórdão.")
        r3 = executar_etapa3(r1, r2, texto_acordao, "Gere a minuta.")

        assert isinstance(r3, ResultadoEtapa3)
        assert len(r3.minuta_completa) > 0


class TestEtapa3StructuredJson:
    """Test structured JSON conversion and fallback behavior in Stage 3."""

    def test_resultado_from_json_payload(self) -> None:
        payload = {
            "minuta_completa": MINUTA_COMPLETA,
            "decisao": "INADMITIDO",
            "fundamentos_decisao": ["Óbice sumular"],
            "itens_evidencia_usados": ["Tema 1/obices_sumulas: Súmula 7 (p.1)"],
        }
        resultado = _resultado_etapa3_from_json(payload)
        assert resultado.decisao == Decisao.INADMITIDO
        assert "Seção III" in resultado.minuta_completa
        assert resultado.fundamentos_decisao == ["Óbice sumular"]
        assert len(resultado.itens_evidencia_usados) == 1

    def test_resultado_from_json_payload_inconclusivo(self) -> None:
        payload = {
            "minuta_completa": "AVISO: Decisão jurídica inconclusiva: Requer análise adicional.",
            "decisao": "INCONCLUSIVO",
            "fundamentos_decisao": ["Dados insuficientes."],
            "itens_evidencia_usados": ["Etapa 1/numero_processo: Processo nº 123 (p.1)"],
        }
        resultado = _resultado_etapa3_from_json(payload)
        assert resultado.decisao == Decisao.INCONCLUSIVO
        assert resultado.aviso_inconclusivo is True
        assert resultado.motivo_bloqueio_codigo == "E3_INCONCLUSIVO"
        assert resultado.motivo_bloqueio_descricao

    def test_executar_etapa3_structured_success(self, monkeypatch) -> None:
        payload = {
            "minuta_completa": MINUTA_COMPLETA,
            "decisao": "INADMITIDO",
            "fundamentos_decisao": ["Óbice sumular"],
            "itens_evidencia_usados": ["Tema 1/obices_sumulas: Súmula 7 (p.1)"],
        }
        monkeypatch.setattr("src.etapa3.chamar_llm_json", lambda **kwargs: payload)

        def _should_not_call_llm(**kwargs):
            raise AssertionError("legacy fallback should not be called")

        monkeypatch.setattr("src.etapa3.chamar_llm", _should_not_call_llm)

        r1 = ResultadoEtapa1(
            numero_processo="123",
            recorrente="EMPRESA ABC LTDA",
            evidencias_campos={
                "numero_processo": CampoEvidencia(
                    citacao_literal="Processo nº 123",
                    pagina=1,
                    ancora="Processo nº 123",
                )
            },
        )
        r2 = ResultadoEtapa2(
            temas=[
                TemaEtapa2(
                    materia_controvertida="Responsabilidade civil",
                    conclusao_fundamentos="Sem dano moral",
                    obices_sumulas=["Súmula 7"],
                    trecho_transcricao="Não restou demonstrada a ocorrência do dano alegado.",
                    evidencias_campos={
                        "obices_sumulas": CampoEvidencia(
                            citacao_literal="Incide a Súmula 7",
                            pagina=1,
                            ancora="Súmula 7",
                        )
                    },
                )
            ]
        )
        resultado = executar_etapa3(
            resultado_etapa1=r1,
            resultado_etapa2=r2,
            texto_acordao="Não restou demonstrada a ocorrência do dano alegado pelo recorrente.",
            prompt_sistema="prompt",
            modelo_override="gpt-4o",
        )
        assert resultado.decisao == Decisao.INADMITIDO
        assert resultado.fundamentos_decisao
        assert resultado.itens_evidencia_usados

    def test_executar_etapa3_structured_failure_fallback(self, monkeypatch) -> None:
        calls = {"json": 0}

        def _fail_json(**kwargs):
            calls["json"] += 1
            raise RuntimeError("json inválido")

        monkeypatch.setattr("src.etapa3.chamar_llm_json", _fail_json)

        class _FakeResponse:
            content = MINUTA_COMPLETA
            tokens = type("T", (), {"total_tokens": 200, "prompt_tokens": 120, "completion_tokens": 80})()

        monkeypatch.setattr("src.etapa3.chamar_llm", lambda **kwargs: _FakeResponse())

        r1 = ResultadoEtapa1(numero_processo="123", recorrente="EMPRESA ABC LTDA")
        r2 = ResultadoEtapa2(temas=[TemaEtapa2(materia_controvertida="Responsabilidade civil", conclusao_fundamentos="Sem dano moral")])
        resultado = executar_etapa3(
            resultado_etapa1=r1,
            resultado_etapa2=r2,
            texto_acordao="Não restou demonstrada a ocorrência do dano alegado pelo recorrente.",
            prompt_sistema="prompt",
            modelo_override="gpt-4o",
        )
        assert calls["json"] == 2
        assert resultado.decisao == Decisao.INADMITIDO

    def test_executar_etapa3_decisao_deterministica_sobrepoe_minuta(self, monkeypatch) -> None:
        payload = {"minuta_completa": MINUTA_COMPLETA, "decisao": "INADMITIDO"}
        monkeypatch.setattr("src.etapa3.chamar_llm_json", lambda **kwargs: payload)

        def _should_not_call_llm(**kwargs):
            raise AssertionError("legacy fallback should not be called")

        monkeypatch.setattr("src.etapa3.chamar_llm", _should_not_call_llm)

        r1 = ResultadoEtapa1(
            numero_processo="123",
            recorrente="EMPRESA ABC LTDA",
            permissivo_constitucional="art. 105, III, a, CF",
            dispositivos_violados=["art. 927 do CC"],
        )
        r2 = ResultadoEtapa2(
            temas=[
                TemaEtapa2(
                    materia_controvertida="Responsabilidade civil",
                    conclusao_fundamentos="Sem óbice identificado para conhecimento do recurso.",
                    obices_sumulas=[],
                )
            ]
        )
        resultado = executar_etapa3(
            resultado_etapa1=r1,
            resultado_etapa2=r2,
            texto_acordao="Não restou demonstrada a ocorrência do dano alegado pelo recorrente.",
            prompt_sistema="prompt",
            modelo_override="gpt-4o",
        )
        assert resultado.decisao == Decisao.ADMITIDO

    def test_executar_etapa3_inconclusivo_forca_aviso(self, monkeypatch) -> None:
        payload = {
            "minuta_completa": "Seção I – Relatório\nSeção II – Análise\nSeção III – Decisão",
            "decisao": "INCONCLUSIVO",
            "fundamentos_decisao": ["Dados conflitantes."],
            "itens_evidencia_usados": [],
        }
        monkeypatch.setattr("src.etapa3.chamar_llm_json", lambda **kwargs: payload)

        def _should_not_call_llm(**kwargs):
            raise AssertionError("legacy fallback should not be called")

        monkeypatch.setattr("src.etapa3.chamar_llm", _should_not_call_llm)
        monkeypatch.setattr(
            "src.etapa3._decidir_admissibilidade_deterministica",
            lambda r1, r2: (Decisao.INCONCLUSIVO, ["Dados insuficientes para conclusão segura."]),
        )

        r1 = ResultadoEtapa1(numero_processo="123", recorrente="EMPRESA ABC LTDA")
        r2 = ResultadoEtapa2(
            temas=[TemaEtapa2(materia_controvertida="Tema X", conclusao_fundamentos="")],
        )
        resultado = executar_etapa3(
            resultado_etapa1=r1,
            resultado_etapa2=r2,
            texto_acordao="Texto curto sem dados suficientes.",
            prompt_sistema="prompt",
            modelo_override="gpt-4o",
        )
        assert resultado.decisao == Decisao.INCONCLUSIVO
        assert resultado.aviso_inconclusivo is True
        assert "AVISO:" in resultado.minuta_completa
        assert resultado.motivo_bloqueio_codigo == "E3_INCONCLUSIVO"
        assert resultado.motivo_bloqueio_descricao


class TestDecisaoDeterministicaEtapa3:
    """Test deterministic admissibility engine used by Stage 3."""

    def test_inadmite_por_sumula_forte(self) -> None:
        r1 = ResultadoEtapa1(
            permissivo_constitucional="art. 105, III, a, CF",
            dispositivos_violados=["art. 927 do CC"],
        )
        r2 = ResultadoEtapa2(
            temas=[TemaEtapa2(conclusao_fundamentos="Tema com óbice", obices_sumulas=["Súmula 7/STJ"])]
        )
        decisao, fundamentos = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao == Decisao.INADMITIDO
        assert any("Súmula 7" in f for f in fundamentos)
        assert any("Regra de precedência aplicada" in f for f in fundamentos)

    def test_inadmite_sem_lastro_minimo_etapa1(self) -> None:
        r1 = ResultadoEtapa1(permissivo_constitucional="", dispositivos_violados=[])
        r2 = ResultadoEtapa2(
            temas=[TemaEtapa2(conclusao_fundamentos="Sem óbice detectado", obices_sumulas=[])]
        )
        decisao, fundamentos = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao == Decisao.INADMITIDO
        assert any("Etapa 1 sem permissivo constitucional" in f for f in fundamentos)

    def test_admite_sem_obice_forte_com_lastro_minimo(self) -> None:
        r1 = ResultadoEtapa1(
            permissivo_constitucional="art. 105, III, a, CF",
            dispositivos_violados=["art. 489 do CPC"],
        )
        r2 = ResultadoEtapa2(
            temas=[TemaEtapa2(conclusao_fundamentos="Sem óbice processual detectado", obices_sumulas=[])]
        )
        decisao, fundamentos = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao == Decisao.ADMITIDO
        assert any("Sem óbice forte identificado" in f for f in fundamentos)

    def test_inconclusivo_sem_conclusoes_etapa2(self) -> None:
        r1 = ResultadoEtapa1(
            permissivo_constitucional="art. 105, III, a, CF",
            dispositivos_violados=["art. 489 do CPC"],
        )
        r2 = ResultadoEtapa2(
            temas=[TemaEtapa2(materia_controvertida="Tema sem conclusão", conclusao_fundamentos="")],
        )
        decisao, fundamentos = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao == Decisao.INCONCLUSIVO
        assert any("sem conclusões/fundamentos" in f.lower() for f in fundamentos)

    def test_precedencia_obice_forte_sobre_indicio_admissibilidade(self) -> None:
        r1 = ResultadoEtapa1(
            permissivo_constitucional="art. 105, III, a, CF",
            dispositivos_violados=["art. 489 do CPC"],
        )
        r2 = ResultadoEtapa2(
            temas=[
                TemaEtapa2(
                    conclusao_fundamentos="Recurso conhecido, porém incide óbice de reexame.",
                    obices_sumulas=["Súmula 7/STJ"],
                )
            ]
        )
        decisao, fundamentos = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao == Decisao.INADMITIDO
        assert any("Regra de precedência aplicada" in f for f in fundamentos)
        assert any("Súmula 7" in f for f in fundamentos)

    def test_conflito_mesmo_nivel_gera_inconclusivo(self) -> None:
        r1 = ResultadoEtapa1(
            permissivo_constitucional="art. 105, III, a, CF",
            dispositivos_violados=["art. 489 do CPC"],
        )
        r2 = ResultadoEtapa2(
            temas=[
                TemaEtapa2(
                    conclusao_fundamentos=(
                        "Há juízo positivo de admissibilidade, mas também não conhecimento do recurso."
                    ),
                    obices_sumulas=[],
                )
            ]
        )
        decisao, fundamentos = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao == Decisao.INCONCLUSIVO
        assert any("Conflito de evidências no mesmo nível" in f for f in fundamentos)


class TestMutationSensitiveDecisionRules:
    """Mutation-sensitive tests for critical deterministic decision rules."""

    def _base_inputs(self) -> tuple[ResultadoEtapa1, ResultadoEtapa2]:
        r1 = ResultadoEtapa1(
            permissivo_constitucional="art. 105, III, a, CF",
            dispositivos_violados=["art. 489 do CPC"],
        )
        r2 = ResultadoEtapa2(
            temas=[
                TemaEtapa2(
                    materia_controvertida="Tema base",
                    conclusao_fundamentos="Sem óbice processual detectado.",
                    obices_sumulas=[],
                    trecho_transcricao="Trecho base.",
                )
            ]
        )
        return r1, r2

    def test_mutation_add_strong_obice_flips_admitido_to_inadmitido(self) -> None:
        r1, r2 = self._base_inputs()
        decisao_base, _ = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao_base == Decisao.ADMITIDO

        r2.temas[0].obices_sumulas = ["Súmula 7/STJ"]
        decisao_mutada, fundamentos_mutados = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao_mutada == Decisao.INADMITIDO
        assert any("Súmula 7" in f for f in fundamentos_mutados)

    def test_mutation_remove_minimum_support_flips_admitido_to_inadmitido(self) -> None:
        r1, r2 = self._base_inputs()
        decisao_base, _ = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao_base == Decisao.ADMITIDO

        r1.permissivo_constitucional = ""
        r1.dispositivos_violados = []
        decisao_mutada, fundamentos_mutados = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao_mutada == Decisao.INADMITIDO
        assert any("Etapa 1 sem permissivo constitucional" in f for f in fundamentos_mutados)

    def test_mutation_insert_conflicting_markers_flips_to_inconclusivo(self) -> None:
        r1, r2 = self._base_inputs()
        decisao_base, _ = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao_base == Decisao.ADMITIDO

        r2.temas[0].conclusao_fundamentos = (
            "Há juízo positivo de admissibilidade, mas também não conhecimento do recurso."
        )
        decisao_mutada, fundamentos_mutados = _decidir_admissibilidade_deterministica(r1, r2)
        assert decisao_mutada == Decisao.INCONCLUSIVO
        assert any("Conflito de evidências no mesmo nível" in f for f in fundamentos_mutados)
